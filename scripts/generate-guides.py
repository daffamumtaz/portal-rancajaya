from __future__ import annotations

import html
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT.parent / "Referensi Dokumen Operasional" / "Panduan-Pengelola-CMS.docx"
DOC_DIR = ROOT / "Dokumen Operasional Pengurus"
PUBLIC_FILES = ROOT / "public" / "admin" / "panduan" / "files"
ASSET_DIR = ROOT / "public" / "admin" / "panduan" / "assets"

GREEN = "174C37"
GREEN_BRIGHT = "166534"
GOLD = "B7791F"
INK = "252B28"
MUTED = "687386"
PALE_GREEN = "E9F4EE"
PALE_GOLD = "FBF3DD"
PALE_RED = "FBEDE7"
BORDER = "D9CDB4"
REPOSITORY_URL = "https://github.com/portalrancajaya/portalrancajaya.github.io"
SUPPORT_URL = "https://wa.me/6289663580475"
SUPPORT_LABEL = "Hubungi Daffa"


def p(text): return {"type": "p", "text": text}
def h(text): return {"type": "h", "text": text}
def steps(*items): return {"type": "steps", "items": list(items)}
def bullets(*items): return {"type": "bullets", "items": list(items)}
def note(title, text, tone="green"): return {"type": "note", "title": title, "text": text, "tone": tone}
def table(headers, rows): return {"type": "table", "headers": headers, "rows": rows}
def image(name, caption): return {"type": "image", "name": name, "caption": caption}
def code(text): return {"type": "code", "text": text}
def link(label, url): return {"type": "link", "label": label, "url": url}


OPERATIONAL = {
    "slug": "panduan-operasional",
    "filename": "Panduan-Operasional-Portal-Rancajaya",
    "eyebrow": "PANDUAN PENGGUNA",
    "title": "Panduan Operasional",
    "subtitle": "Halaman Pengelolaan Website",
    "portal": "PORTAL DESA RANCAJAYA",
    "description": "Panduan langkah demi langkah untuk memperbarui informasi desa melalui halaman admin.",
    "audience": "Warga dan Perangkat Desa yang Bertugas Mengelola Website Rancajaya",
    "sections": [
        ("Mengenal Halaman Pengelolaan", [
            p("Halaman pengelolaan membantu perangkat desa memperbarui isi Portal Rancajaya melalui browser. Perubahan yang disimpan akan diteruskan ke website secara otomatis."),
            link("Bantuan teknis Daffa melalui WhatsApp", SUPPORT_URL),
            steps("Masuk ke halaman admin.", "Pilih jenis informasi yang ingin diperbarui.", "Isi atau perbaiki data.", "Klik Terbitkan, lalu tunggu beberapa menit sampai website berubah."),
            table(["Dapat ditambah dan diubah", "Hanya memperbarui data yang tersedia"], [["Berita, APBDes, statistik potensi, kelompok tani, produksi pertanian, UMKM, layanan, kontak, dan galeri", "Struktur organisasi, dusun, pengaturan umum, serta halaman website"]]),
            image("admin-dashboard.png", "Ilustrasi susunan menu berdasarkan konfigurasi Portal Rancajaya."),
            note("Penting", "Menu tidak menyediakan penghapusan permanen. Informasi yang tidak ingin ditampilkan dapat diubah menjadi Draf atau dimatikan melalui pilihan Tampilkan di Publik.", "gold"),
        ]),
        ("Masuk ke Halaman Admin", [
            p("Buka alamat https://portalrancajaya.github.io/admin melalui Chrome, Edge, atau Firefox."),
            steps("Klik Masuk dengan GitHub.", "Masukkan akun GitHub yang telah diberi akses oleh pengembang.", "Jika diminta, setujui izin akses.", "Setelah berhasil, daftar menu pengelolaan akan tampil."),
            note("Keamanan akun", "Jangan menuliskan kata sandi di dokumen, grup percakapan, atau catatan yang dapat dibaca umum. Aktifkan verifikasi dua langkah pada akun GitHub.", "gold"),
        ]),
        ("Berita dan Informasi", [
            p("Menu Berita & Informasi dipakai untuk Berita, Agenda, Artikel, dan Pengumuman."),
            steps("Pilih Berita & Informasi.", "Klik + Informasi.", "Isi judul, tanggal, kategori, ringkasan, penulis, dan konten.", "Tambahkan thumbnail jika tersedia.", "Pilih Draf jika belum boleh tampil atau Terbit jika sudah siap, lalu klik Terbitkan."),
            image("admin-berita.png", "Contoh bidang utama saat menulis informasi baru."),
            bullets("Tampilkan di Beranda digunakan untuk informasi yang perlu disorot.", "Tautan Artikel Lengkap dan Nama Media Sumber diisi jika tulisan lengkap berada di media lain.", "Judul halaman dibuat otomatis dari judul berita; pengelola tidak perlu menulis alamat halaman sendiri."),
        ]),
        ("Profil Desa dan Dusun", [
            p("Data identitas desa berada pada Pengaturan dan Halaman Website. Data tiga dusun berada pada menu Dusun."),
            steps("Buka Pengaturan > Pengaturan Umum untuk visi, misi, tujuan, sejarah, alamat, kontak, jam layanan, media sosial, peta, dan tahun APBDes aktif.", "Buka Halaman Website > Halaman Profil untuk sambutan, nama kepala desa, luas wilayah, penduduk, batas wilayah, dan demografi.", "Buka Dusun untuk memperbarui kepala dusun, penduduk, kepala keluarga, luas wilayah, dan deskripsi."),
            note("Sumber data", "Gunakan dokumen resmi desa. Kolom tahun data, sumber data, dan status validasi membantu pembaca mengetahui asal dan kepastian informasi.", "green"),
        ]),
        ("Struktur Organisasi", [
            p("Menu Struktur Organisasi hanya untuk memperbarui posisi yang sudah disiapkan. Hubungan atasan menentukan garis susunan organisasi di halaman Profil."),
            steps("Pilih posisi yang akan diperbarui.", "Periksa nama dan jabatan.", "Pilih Atasan Langsung yang benar. Kepala Desa tidak memiliki atasan.", "Atur urutan agar susunan rapi.", "Isi deskripsi, tahun data, sumber, dan status validasi bila tersedia, lalu Terbitkan."),
            note("Jangan menambah posisi dari panel", "Jika ada perubahan susunan yang memerlukan posisi baru atau penghapusan posisi, hubungi pengembang agar garis organisasi dan berkas sumber tetap benar.", "gold"),
        ]),
        ("Transparansi APBDes", [
            p("Setiap baris APBDes mewakili satu sumber Pendapatan atau satu pos Belanja untuk tahun dan versi tertentu."),
            steps("Pilih Transparansi APBDes lalu klik + Pos APBDes.", "Isi nama sumber atau pos.", "Pilih Pendapatan atau Belanja.", "Isi tahun serta versi Awal atau Perubahan.", "Isi jumlah rupiah dengan angka tanpa simbol Rp dan tanpa titik.", "Lengkapi sumber serta status validasi, kemudian Terbitkan."),
            image("admin-apbdes.png", "Bidang utama untuk satu pos APBDes."),
            note("Ketelitian angka", "Satu pos hanya boleh dicatat sekali pada tahun dan versi yang sama. Gunakan angka dari dokumen APBDes resmi yang sudah boleh dipublikasikan.", "gold"),
        ]),
        ("Potensi dan Kelompok Tani", [
            p("Informasi pertanian dikelola melalui tiga menu: Statistik Potensi, Kelompok Tani, dan Produksi Pertanian."),
            table(["Menu", "Isi yang dikelola"], [["Statistik Potensi", "Angka ringkas pertanian, peternakan, dan sarana pertanian"], ["Kelompok Tani", "Nama kelompok, ketua, dusun, komoditas, luas, dan anggota"], ["Produksi Pertanian", "Komoditas, luas, produktivitas, produksi, dan catatan"]]),
            steps("Pilih menu yang sesuai.", "Tambah data baru atau buka data yang tersedia.", "Gunakan satuan yang konsisten.", "Cantumkan tahun dan sumber data, misalnya Programa BPP Patokbeusi.", "Terbitkan setelah isian selesai."),
        ]),
        ("UMKM Desa", [
            steps("Pilih UMKM lalu klik + UMKM.", "Isi nama UMKM, pemilik, kategori, produk utama, dusun, kontak, foto, dan deskripsi.", "Nomor kontak boleh diawali 08; website akan mengarahkannya ke WhatsApp jika formatnya sesuai.", "Pilih Tampilkan sebagai Unggulan hanya untuk UMKM yang perlu disorot.", "Terbitkan."),
            note("Foto dan deskripsi", "Gunakan foto produk milik UMKM dan deskripsi singkat yang menjelaskan apa yang dijual. Jangan mencantumkan alamat rumah yang tidak disetujui pemilik.", "green"),
        ]),
        ("Layanan Surat dan Kontak", [
            p("Administrasi Daring menampilkan jenis surat, persyaratan, dan tautan formulir. Kontak menampilkan nomor layanan yang dapat dipilih warga."),
            steps("Buka Administrasi Daring untuk menambah atau memperbarui nama layanan, kategori, deskripsi, tautan formulir, persyaratan, dan urutan.", "Gunakan tautan formulir resmi: https://forms.gle/jmu5QRToacBfipTi6.", "Buka Kontak untuk memperbarui nama, jenis, telepon, alamat, deskripsi, dan urutan.", "Ubah menjadi Draf atau matikan Tampilkan di Publik jika layanan sementara tidak boleh terlihat."),
            note("Data warga", "Portal hanya menampilkan informasi dan tautan formulir. Jangan mengunggah KTP, KK, NIK, hasil formulir, atau surat pribadi ke halaman publik.", "red"),
        ]),
        ("Galeri Foto", [
            steps("Pilih Galeri Foto lalu klik + Foto Galeri.", "Isi judul, tanggal, kategori, foto, teks alternatif, dan deskripsi.", "Teks alternatif menjelaskan isi foto secara singkat, misalnya Warga mengikuti lokakarya di Balai Desa Rancajaya.", "Terbitkan."),
            image("admin-media.png", "Ilustrasi halaman pemilihan dan pengunggahan gambar."),
            bullets("Gunakan JPG, PNG, atau WebP.", "Usahakan ukuran foto tidak lebih dari 5 MB; saat ini batas ukuran belum dipaksakan oleh halaman admin.", "Gunakan nama berkas yang jelas dan tanpa karakter aneh.", "Pilih foto terang, tidak buram, dan tidak memperlihatkan dokumen pribadi warga."),
        ]),
        ("Beranda dan Pengaturan Umum", [
            p("Halaman Website berisi teks dan gambar khusus untuk Beranda, Profil, Potensi, dan Layanan. Pengaturan berisi identitas yang dipakai di banyak halaman."),
            steps("Buka Halaman Website > Halaman Beranda untuk judul utama, subjudul, gambar, sorotan, dan statistik.", "Buka Halaman Potensi untuk narasi pertanian, Gapoktan, peternakan domba, dan galeri peternakan.", "Buka Halaman Layanan untuk judul serta subjudul halaman layanan.", "Buka Pengaturan Umum untuk alamat kantor, kontak, media sosial, peta, dan tahun APBDes aktif."),
            note("Dampak perubahan", "Perubahan Pengaturan Umum dapat muncul di beberapa halaman sekaligus, termasuk header, footer, Profil, dan Layanan.", "gold"),
        ]),
        ("Menyimpan dan Mengatasi Masalah", [
            h("Pilihan penayangan"),
            table(["Pilihan", "Hasil"], [["Draf", "Informasi disimpan tetapi tidak ditampilkan di website"], ["Terbit", "Informasi dapat tampil setelah website selesai diperbarui"], ["Tampilkan di Publik dimatikan", "Informasi disembunyikan tanpa menghapus berkasnya"]]),
            h("Masalah umum"),
            bullets("Perubahan belum terlihat: tunggu beberapa menit lalu muat ulang halaman.", "Tidak bisa masuk: pastikan akun GitHub memiliki akses tulis ke repository Portal Rancajaya.", "Foto tidak muncul: unggah ulang JPG, PNG, atau WebP dengan nama berkas sederhana.", "Website gagal diperbarui: catat informasi terakhir yang diubah dan hubungi PIC teknis.", "Salah isi: buka kembali data, perbaiki, lalu Terbitkan ulang."),
            note("Bantuan teknis", "Sampaikan halaman yang bermasalah, perubahan terakhir, waktu kejadian, dan tangkapan layar pesan kesalahan kepada Daffa melalui WhatsApp.", "green"),
        ]),
    ],
}


DEVELOPMENT = {
    "slug": "panduan-pengembangan",
    "filename": "Panduan-Pengembangan-Portal-Rancajaya",
    "eyebrow": "PANDUAN TEKNIS",
    "title": "Panduan Pengembangan",
    "subtitle": "Arsitektur, Ekstensi, dan Pemeliharaan Portal",
    "portal": "PORTAL DESA RANCAJAYA",
    "description": "Panduan teknis untuk memahami, memperluas, memelihara, dan memulihkan Portal Desa Rancajaya.",
    "audience": "Pengembang dan PIC Teknis Portal Rancajaya",
    "sections": [
        ("Arsitektur dan Alur Publikasi", [
            p("Portal Rancajaya adalah website statis berbasis Astro. Konten dikelola sebagai berkas Markdown dan JSON di GitHub; tidak ada server aplikasi atau database produksi."),
            image("arsitektur-portal.png", "Alur publikasi Portal Rancajaya dari pengelola sampai pengunjung."),
            table(["Tahap", "Tanggung jawab"], [["Decap CMS", "Menulis perubahan konten sebagai commit Git"], ["GitHub", "Menyimpan kode, konten, aset, dan riwayat"], ["GitHub Actions", "Memasang dependency, memvalidasi konten, dan membangun dist"], ["GitHub Pages", "Menyajikan hasil statis kepada pengunjung"], ["Netlify", "Hanya meneruskan OAuth untuk proses masuk GitHub"]]),
            note("Batas sistem", "Portal bukan database kependudukan, sistem persuratan otomatis, atau tempat menyimpan KTP, KK, NIK, respons formulir, dan surat pribadi warga.", "gold"),
        ]),
        ("Stack dan Alasan Pemilihan", [
            table(["Komponen", "Peran dan alasan"], [["Astro 5", "Routing berbasis berkas dan keluaran statis; sesuai untuk portal informasi dengan JavaScript minimum"], ["Tailwind CSS 3", "Token desain dan kelas utilitas yang konsisten tanpa framework UI tambahan"], ["Content Collections + Zod", "Validasi struktur Markdown/JSON saat build sebelum perubahan tayang"], ["Decap CMS 3", "Antarmuka pengelolaan berbasis Git tanpa database CMS"], ["Chart.js 4", "Grafik APBDes dan demografi dari data yang sudah tersedia di halaman"], ["GitHub Actions + Pages", "Build dan hosting produksi gratis dalam satu riwayat repository"], ["Netlify", "OAuth proxy saja; bukan hosting produksi"]]),
            link("Repository resmi", REPOSITORY_URL),
            p("Situs produksi: https://portalrancajaya.github.io."),
            note("Tidak dipakai", "Portal tidak memakai Pagefind, Leaflet, Formspree, Vitest, atau backend database. Jangan menambahkan dependency hanya karena tercantum pada proyek referensi.", "gold"),
        ]),
        ("Struktur Repository", [
            code("src/pages/                 rute: beranda, profil, informasi, potensi, layanan\nsrc/components/            Navbar, Footer, Seo, kartu, dan Hero\nsrc/layouts/Layout.astro   shell halaman dan metadata global\nsrc/content/               koleksi Markdown serta data JSON\nsrc/content/config.ts      skema Zod seluruh konten\nsrc/styles/global.css      kelas komponen dan gaya dasar\nsrc/utils/content.ts       publikasi, URL aset, WhatsApp, formulir, peta\npublic/admin/config.yml    collection dan field Decap CMS\npublic/admin/index.html    halaman masuk dan akses panduan\npublic/images/             gambar publik dan hasil unggahan\npublic/vendor/             Chart.js yang dilayani lokal\n.github/workflows/         build dan deploy GitHub Pages"),
            note("Sumber kebenaran", "Untuk setiap field, src/content/config.ts, public/admin/config.yml, contoh konten, dan halaman yang membaca data harus tetap konsisten.", "red"),
        ]),
        ("Alur Data dan Model Konten", [
            p("Konten collection disimpan sebagai Markdown, sedangkan pengaturan dan teks halaman tunggal disimpan sebagai JSON. Halaman mengambil data melalui getCollection atau getEntry lalu memfilter status publikasi."),
            table(["Kelompok data", "Sumber", "Konsumen utama"], [["Berita", "src/content/berita/*.md", "Beranda dan halaman Informasi"], ["APBDes", "src/content/apbdes/*.md", "Profil dan grafik riwayat APBDes"], ["Potensi", "potensi-statistik, kelompok-tani, produksi-pertanian", "Halaman Potensi"], ["UMKM dan galeri", "src/content/umkm serta galeri", "Beranda, Potensi, dan Informasi visual"], ["Struktur dan dusun", "struktur-organisasi serta dusun", "Halaman Profil"], ["Layanan dan kontak", "layanan-administrasi serta kontak", "Beranda dan Layanan"], ["Pengaturan", "pengaturan/umum.json", "Navbar, Footer, peta, dan identitas desa"], ["Halaman", "halaman/*.json", "Teks khusus Beranda, Profil, Potensi, dan Layanan"]]),
            table(["Status", "Perilaku"], [["release_status: Draf", "isPublishedContent mengembalikan false"], ["release_status: Terbit", "konten boleh ditampilkan"], ["release_status: Siap", "nilai lama yang tetap dianggap tampil"], ["aktif: false", "konten selalu disembunyikan"]]),
            note("Validasi bukan verifikasi", "Zod memeriksa tipe dan struktur data, bukan kebenaran faktual. Tahun, sumber, dan status validasi tetap harus diperiksa terhadap dokumen resmi.", "gold"),
        ]),
        ("Menjalankan Proyek Secara Lokal", [
            steps("Gunakan Node.js 20 atau versi kompatibel.", "Clone repository dan masuk ke folder portalrancajaya.github.io.", "Jalankan npm ci agar dependency mengikuti package-lock.json.", "Jalankan npm run dev.", "Buka http://localhost:4321/.", "Jalankan npm run build dan npm run preview sebelum commit."),
            code("git clone https://github.com/portalrancajaya/portalrancajaya.github.io.git\ncd portalrancajaya.github.io\nnpm ci\nnpm run dev\nnpm run build\nnpm run preview"),
            note("Batas pengujian", "Belum ada automated test di package.json. npm run build adalah pemeriksaan minimum untuk schema, import, dan kompilasi; tampilan utama dan admin tetap harus diperiksa di browser.", "gold"),
        ]),
        ("Menambah Halaman Baru", [
            steps("Buat file .astro di src/pages; nama folder atau file menentukan URL.", "Gunakan Layout dari src/layouts/Layout.astro.", "Isi title dan description lokal serta schema JSON-LD bila relevan.", "Gunakan getCollection/getEntry untuk data dan isPublishedContent untuk collection publik.", "Gunakan resolvePublicPath untuk gambar dari konten.", "Tambahkan navigasi di src/components/Navbar.astro dan Footer.astro bila halaman memang publik.", "Uji desktop, ponsel, tautan aktif, base path, lalu jalankan npm run build."),
            code("---\nimport Layout from '../layouts/Layout.astro';\nconst base = import.meta.env.BASE_URL.replace(/\\/$/, '');\n---\n<Layout title=\"Judul Halaman\" description=\"Deskripsi lokal\">\n  <main class=\"section\">...</main>\n</Layout>"),
            note("Navigasi", "Navbar.astro dan Footer.astro mempunyai daftar tautan masing-masing. Jika menu baru harus tampil di keduanya, perbarui kedua file secara bersamaan.", "green"),
        ]),
        ("Menambah atau Mengubah Collection", [
            steps("Definisikan atau ubah schema Zod di src/content/config.ts.", "Daftarkan collection pada export collections di file yang sama.", "Buat folder konten dan satu contoh berkas yang valid.", "Tambahkan collection serta field yang identik di public/admin/config.yml.", "Atur create, delete, identifier_field, summary, dan slug sesuai kebutuhan.", "Sesuaikan komponen atau halaman yang membaca data.", "Jalankan npm run build, lalu uji tambah dan edit melalui admin."),
            code("const contohCollection = defineCollection({\n  type: 'content',\n  schema: z.object({\n    nama: z.string(),\n    urutan: z.number(),\n    ...publicationFields,\n  }),\n});"),
            note("Relation field", "identifier_field, search_fields, value_field, nama folder, nama collection, dan nilai frontmatter harus konsisten. Perubahan salah satu tanpa yang lain dapat merusak pilihan relasi atau build.", "gold"),
        ]),
        ("Menambah Visualisasi Data", [
            p("Padanan fitur infografis pada Portal Rancajaya adalah grafik demografi dan APBDes di src/pages/profil.astro. Chart.js dilayani lokal dari public/vendor/chart.umd.js."),
            steps("Tentukan sumber data dan satuan yang akan divisualisasikan.", "Validasi field sumber di src/content/config.ts dan CMS jika datanya dapat diedit.", "Ubah data menjadi array label/value di frontmatter halaman.", "Tambahkan canvas dengan id unik, role img, dan aria-label yang menjelaskan grafik.", "Muat chartScriptUrl satu kali dan buat grafik setelah DOMContentLoaded.", "Sediakan angka ringkas, legenda teks, atau tabel agar informasi tidak hanya bergantung pada grafik.", "Uji kondisi data kosong, banyak kategori, layar kecil, dan pergantian periode."),
            code("const chartScriptUrl = `${import.meta.env.BASE_URL.replace(/\\/$/, '')}/vendor/chart.umd.js`;\n\n<canvas id=\"chartBaru\" role=\"img\" aria-label=\"Deskripsi grafik\"></canvas>\n<script is:inline src={chartScriptUrl}></script>"),
            note("Aksesibilitas data", "Canvas tidak menggantikan teks. Nilai penting harus tetap tersedia sebagai ringkasan, legenda, atau tabel yang dapat dibaca tanpa menjalankan JavaScript.", "green"),
        ]),
        ("Mengubah Design System", [
            p("Token utama berada di tailwind.config.mjs, sedangkan pola komponen berulang berada di src/styles/global.css. Perubahan token harus lebih dahulu dipilih daripada mengganti warna satu per satu di halaman."),
            table(["Bagian", "Lokasi dan isi"], [["Warna", "primary, gold, ink, muted, paper, dan line di tailwind.config.mjs"], ["Font", "Lexend untuk display; Plus Jakarta Sans/Inter untuk body"], ["Radius", "2xl, xl, lg, md, dan pill"], ["Kartu", ".card, .card-hover, .stat-card, .stat-card-dark"], ["Tombol", ".btn-primary, .btn-gold, .btn-outline, .btn-ghost"], ["Label", ".pill, .pill-primary, .pill-gold, .eyebrow"], ["Tata letak", ".container-page, .section, .section-header, .section-title"]]),
            steps("Ubah token atau kelas dasar.", "Cari warna atau ukuran hardcoded yang perlu dinormalisasi.", "Periksa kontras teks, focus state, dan ukuran target sentuh.", "Uji Beranda, Profil, Informasi, Potensi, Layanan, serta admin pada desktop dan ponsel.", "Jalankan npm run build."),
            note("Jangan mengubah merek parsial", "Perubahan warna, font, atau radius harus diterapkan sebagai satu sistem. Hindari menambahkan warna baru pada komponen tunggal tanpa peran desain yang jelas.", "gold"),
        ]),
        ("Konfigurasi CMS dan Publikasi", [
            p("Decap CMS memakai backend GitHub pada branch main dan OAuth proxy Netlify. Editorial Workflow tidak diaktifkan; status Draf/Terbit berasal dari field konten Portal Rancajaya."),
            code("backend:\n  name: github\n  repo: portalrancajaya/portalrancajaya.github.io\n  branch: main\n  base_url: https://api.netlify.com\n  auth_endpoint: auth"),
            table(["Operasi", "Collection"], [["Tambah + edit", "berita, APBDes, statistik potensi, kelompok tani, produksi, UMKM, layanan, kontak, galeri"], ["Edit saja", "struktur organisasi, dusun, pengaturan umum, halaman website"], ["Hapus", "dinonaktifkan pada seluruh collection"]]),
            bullets("Jaga urutan field agar mudah dipakai pengelola.", "Berikan hint dan pola nilai untuk field yang rawan salah.", "Jangan menyimpan rahasia, token, atau data pribadi pada config maupun konten.", "Perubahan config CMS harus diuji bersama schema Zod dan contoh konten."),
        ]),
        ("Aset, Base Path, dan URL", [
            p("Repository sudah memakai GitHub Pages user site: https://portalrancajaya.github.io dengan base '/'. Path aset CMS melalui public_folder harus mengikuti base produksi."),
            bullets("Gunakan resolvePublicPath untuk path gambar dari Markdown atau JSON.", "Gunakan import.meta.env.BASE_URL untuk aset statis yang dirangkai di kode.", "Jangan menambahkan prefix repository pada URL aset atau tautan internal.", "Periksa gambar, favicon, vendor Chart.js, halaman admin, dan tautan unduhan panduan setelah perubahan base.", "Pertahankan nama berkas sederhana, format web, dan ukuran yang wajar."),
            note("Situs GitHub Pages", "Konfigurasi site, base, public_folder, canonical, sitemap, callback OAuth, dan dokumentasi harus tetap mengarah ke portalrancajaya.github.io. Untuk user site, base biasanya '/'.", "gold"),
        ]),
        ("SEO dan Data Terstruktur", [
            p("Seo.astro dan Layout.astro menjadi pusat title, description, canonical URL, Open Graph, favicon, sitemap, serta JSON-LD profil desa."),
            bullets("Berikan title dan description berbeda pada setiap halaman.", "Gunakan kata kunci lokal secara alami: Desa Rancajaya, Patokbeusi, Subang, Jawa Barat.", "Pastikan canonical dan sitemap memakai alamat produksi yang benar.", "robots.txt harus mengizinkan halaman publik dan tidak perlu mengindeks /admin.", "JSON-LD hanya memuat identitas, alamat, dan koordinat yang telah diverifikasi.", "Gunakan alt text yang menjelaskan isi gambar, bukan nama file."),
            note("Saran domain resmi", "Saat ini portal masih memakai domain GitHub Pages, belum domain desa.id. Bila desa ingin menggunakan domain resmi, rencanakan perubahan DNS, canonical, sitemap, OAuth callback, dan pengujian URL sebagai satu pekerjaan terjadwal; jangan mengganti alamat sebagian-sebagian.", "green"),
        ]),
        ("Deployment dan CI/CD", [
            p("Workflow .github/workflows/deploy.yml berjalan pada push ke main atau workflow_dispatch."),
            code("checkout -> setup-node 20 -> npm ci -> npm run build\n-> upload-pages-artifact ./dist -> deploy-pages"),
            steps("Push perubahan ke branch yang dituju.", "Buka tab Actions dan periksa workflow terbaru.", "Jika build gagal, buka job Build with Astro dan baca error pertama yang relevan.", "Perbaiki sumber di repository lalu push commit baru.", "Setelah deploy, uji halaman publik, admin, gambar, sitemap, dan unduhan panduan."),
            note("Netlify bukan hosting produksi", "netlify.toml dipertahankan untuk layanan OAuth. Status website produksi diperiksa di GitHub Actions dan GitHub Pages.", "green"),
        ]),
        ("OAuth dan Akses Admin", [
            p("Pengguna Decap CMS memerlukan akun GitHub dengan akses tulis ke repository. Netlify meneruskan OAuth melalui base_url dan auth_endpoint."),
            bullets("Gunakan akun resmi desa sebagai pemilik utama jika memungkinkan.", "Aktifkan 2FA pada GitHub dan Netlify.", "Pastikan callback OAuth mengikuti alamat admin produksi.", "Cabut akses anggota yang tidak lagi bertugas.", "Jangan menaruh token, password, client secret, atau recovery code di repository maupun panduan.", "Catat pemilik repository, OAuth app, Netlify site, domain, dan kontak pemulihan di penyimpanan internal desa."),
            note("Panduan admin bersifat publik", "Halaman dan berkas panduan berada di public/admin tanpa autentikasi tambahan. Panduan boleh menjelaskan proses, tetapi tidak boleh memuat rahasia atau daftar kredensial.", "red"),
        ]),
        ("Keterbatasan dan Utang Teknis", [
            table(["Keterbatasan", "Dampak dan mitigasi"], [["Rebuild pada setiap commit", "Tidak ada preview situs produksi secara real-time; konten baru tayang setelah GitHub Actions selesai, biasanya sekitar 1-2 menit."], ["Akun GitHub wajib", "CMS memerlukan akun GitHub dengan akses tulis; siapkan pemilik dan admin cadangan."], ["OAuth bergantung pada Netlify", "Login Decap CMS bergantung pada layanan eksternal Netlify; catat pemilik site, callback, dan kontak pemulihan."], ["Tidak ada offline editing", "Perubahan membutuhkan koneksi internet dan akses ke GitHub/Netlify."], ["Unggahan gambar belum dibatasi otomatis", "CMS belum menolak ukuran file besar; 5 MB adalah rekomendasi editorial. Kompres gambar sebelum unggah."], ["Optimasi gambar hanya saat build", "Tidak ada CDN resizing saat pengunjung membuka situs; gunakan format web dan ukuran gambar yang wajar."], ["Bahasa tunggal", "Antarmuka dan konten saat ini berbahasa Indonesia; belum ada fitur multi-language."], ["Belum ada automated test", "npm run build dan uji browser wajib sebelum deploy."], ["Data faktual tidak terverifikasi otomatis", "Gunakan sumber_data, tahun_data, dan status_validasi."], ["Formulir surat berada di Google Forms", "Portal hanya menautkan formulir dan tidak mengendalikan retensi respons."], ["Belum memakai domain desa.id", "Situs masih berada di portalrancajaya.github.io; domain resmi desa dapat dipertimbangkan sebagai pengembangan berikutnya."]]),
            note("Pemeliharaan technical debt", "Setiap utang teknis baru harus dicatat bersama dampak, pemilik keputusan, mitigasi sementara, dan syarat kapan perlu diselesaikan. Perbarui panduan pada perubahan yang memengaruhi pengguna atau pengembang.", "gold"),
        ]),
        ("Recovery Playbook", [
            h("Build gagal setelah edit konten"),
            steps("Buka GitHub Actions dan baca error pertama.", "Cari berkas konten serta field yang disebut pada pesan Zod atau import.", "Perbaiki melalui CMS atau GitHub.", "Push perbaikan dan tunggu workflow baru."),
            h("Rollback perubahan yang sudah tayang"),
            code("git log --oneline\ngit show <commit>\ngit revert <commit>\ngit push origin main"),
            h("Pemulihan konfigurasi repository"),
            steps("Pastikan repository resmi adalah portalrancajaya/portalrancajaya.github.io.", "Pastikan GitHub Pages memakai sumber GitHub Actions dan user site portalrancajaya.github.io.", "Periksa backend.repo di public/admin/config.yml.", "Periksa site dan base di astro.config.mjs serta public_folder CMS.", "Periksa OAuth app, callback, dan konfigurasi Netlify.", "Jalankan build lalu uji halaman publik dan login admin."),
            note("Cadangan tambahan", "Simpan aset asli, dokumen sumber APBDes, daftar akses internal, dan konfigurasi domain pada penyimpanan resmi desa di luar repository publik.", "gold"),
        ]),
        ("Diagnostik Masalah Umum", [
            table(["Gejala", "Lokasi diagnosis"], [["CMS tidak dapat masuk", "OAuth Netlify, akses GitHub, callback, dan console browser"], ["Build gagal setelah edit", "GitHub Actions, schema Zod, serta frontmatter/JSON"], ["Gambar 404", "public_folder, base path, resolvePublicPath, dan nama file"], ["Konten Draf tampil", "isPublishedContent dan filter getCollection"], ["Grafik kosong", "data sumber, id canvas, chartScriptUrl, dan waktu render"], ["Sitemap/canonical salah", "astro.config.mjs, Seo.astro, dan hasil dist"], ["Panduan tidak terbuka", "path admin, index.html, base URL, dan hasil dist"], ["Dokumentasi tertinggal", "Bandingkan perubahan terakhir dengan Panduan Operasional, Panduan Pengembangan, dan halaman admin"]]),
            p("Catat URL, waktu kejadian, commit terakhir, pesan error lengkap, browser/perangkat, serta perubahan yang dilakukan sebelum masalah muncul."),
            note("Urutan diagnosis", "Reproduksi masalah, tentukan apakah sumbernya konten atau kode, periksa commit terkait, lakukan perbaikan terkecil, lalu uji build dan halaman yang terdampak.", "green"),
        ]),
        ("Pengembangan Lanjutan", [
            bullets("Pertahankan portal sebagai website informasi statis selama kebutuhan backend belum tervalidasi.", "Jangan menerima upload KTP/KK atau menyimpan respons formulir pada repository publik.", "Jika menambah fitur interaktif, dokumentasikan pemilik data, retensi, akses, risiko, fallback, dan cara pemulihan.", "Tambahkan automated test bila logika transformasi data, filter publikasi, atau normalisasi URL semakin kompleks.", "Update Panduan Operasional, Panduan Pengembangan, halaman admin, dan technical debt setiap kali menu, field, route, URL, dependency, deployment, atau ownership berubah."),
            note("Definition of Done", "Development work is not complete until npm run build passes, affected public and admin pages are checked, no sensitive data is exposed, and the related documentation is updated in the same change. Documentation is part of the deliverable, not an afterthought.", "green"),
        ]),
    ],
}


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    if shd.getparent() is None: tc_pr.append(shd)


def set_cell_margins(cell, top=110, start=130, bottom=110, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
        if node.getparent() is None: tc_mar.append(node)


def set_font(run, name="Calibri", size=11, color=INK, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr(); marker = OxmlElement("w:tblHeader"); marker.set(qn("w:val"), "true"); tr_pr.append(marker)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(twips)); tc_w.set(qn("w:type"), "dxa")
    if tc_w.getparent() is None: tc_pr.append(tc_w)


def add_page_field(paragraph, instruction):
    run = paragraph.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = instruction
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"; end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end): run._r.append(node)
    set_font(run, size=8, color=MUTED)


def clear_document(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"): body.remove(child)


def prepare_doc(doc, guide):
    clear_document(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.left_margin = Inches(.86); section.right_margin = Inches(.86)
    section.top_margin = Inches(.75); section.bottom_margin = Inches(.75)
    section.header_distance = Inches(.3); section.footer_distance = Inches(.35)
    section.different_first_page_header_footer = True
    # The reference document enables separate odd/even headers. Disable that
    # setting so every content page uses the same complete page-number footer.
    doc.settings.odd_and_even_pages_header_footer = False

    style_names = {style.name for style in doc.styles}
    normal = doc.styles["Normal"] if "Normal" in style_names else doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.18
    for style_name, size in (("Heading 1", 24), ("Heading 2", 14), ("Heading 3", 11.5)):
        style = next((item for item in doc.styles if item.name == style_name), None)
        if style is None: style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Cambria"; style.font.size = Pt(size); style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN); style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]; hp.clear(); hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(hp.add_run(f"{guide['title']}  ·  Portal Desa Rancajaya"), size=8, color=MUTED)
    p_pr = hp._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4"); bottom.set(qn("w:color"), BORDER); borders.append(bottom); p_pr.append(borders)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]; fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(fp.add_run("Halaman "), size=8, color=MUTED); add_page_field(fp, "PAGE")
    set_font(fp.add_run(" dari "), size=8, color=MUTED); add_page_field(fp, "NUMPAGES")

    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_footer.paragraphs[0].clear()

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields")) or OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    if update.getparent() is None: settings.append(update)


def bullet_num_id(doc):
    cached = getattr(doc, "_guide_bullet_num_id", None)
    if cached is not None: return cached
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType"); multi.set(qn("w:val"), "singleLevel"); abstract.append(multi)
    level = OxmlElement("w:lvl"); level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet"); level.append(fmt)
    text = OxmlElement("w:lvlText"); text.set(qn("w:val"), "•"); level.append(text)
    ppr = OxmlElement("w:pPr"); tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360"); ppr.append(ind); level.append(ppr); abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id)); aid = OxmlElement("w:abstractNumId"); aid.set(qn("w:val"), str(abstract_id)); num.append(aid); numbering.append(num)
    doc._guide_bullet_num_id = num_id
    return num_id


def add_bullet(doc, text):
    par = doc.add_paragraph(text)
    ppr = par._p.get_or_add_pPr(); numpr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(bullet_num_id(doc)))
    numpr.append(ilvl); numpr.append(numid); ppr.append(numpr)
    return par


def add_hyperlink(paragraph, label, url, size=10.5):
    """Append a styled external hyperlink to a Word paragraph."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), relationship_id)
    run = paragraph.add_run(label); set_font(run, size=size, color="0563C1")
    run._element.get_or_add_rPr().append(OxmlElement("w:u"))
    underline = run._element.rPr.find(qn("w:u")); underline.set(qn("w:val"), "single")
    paragraph._p.remove(run._r); hyperlink.append(run._r); paragraph._p.append(hyperlink)
    return hyperlink


def add_table_separator(doc):
    """Prevent Word from merging adjacent tables that use different column grids."""
    par = doc.add_paragraph()
    par.paragraph_format.space_before = Pt(0)
    par.paragraph_format.space_after = Pt(0)
    par.paragraph_format.line_spacing = Pt(1)
    set_font(par.add_run("\u200b"), size=1)


def add_cover(doc, guide):
    for _ in range(3): doc.add_paragraph()
    ep = doc.add_paragraph(); ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(ep.add_run(guide["eyebrow"]), size=12, color=GOLD, bold=True)
    ep.paragraph_format.space_after = Pt(14)
    for text in (guide["title"], guide["subtitle"]):
        par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(par.add_run(text), "Cambria", 25 if text == guide["title"] else 22, GREEN, True)
        par.paragraph_format.space_after = Pt(3)
    rule = doc.add_paragraph(); rule.paragraph_format.space_before = Pt(9); rule.paragraph_format.space_after = Pt(16)
    p_pr = rule._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "12"); bottom.set(qn("w:color"), GOLD); borders.append(bottom); p_pr.append(borders)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run(guide["portal"]), size=12, color=GREEN, bold=True)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run(guide["description"]), size=10.5, color=MUTED, italic=True)
    doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run("Ditujukan untuk: "), size=10.5, bold=True); set_font(par.add_run(guide["audience"]), size=10.5, bold=True)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run("Bantuan teknis: "), size=10, color=MUTED); add_hyperlink(par, "Daffa via WhatsApp", SUPPORT_URL, size=10)
    doc.add_paragraph()
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run("Disusun oleh"), size=9, color=MUTED)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run("KKN-T IPB 2026"), "Cambria", 13, INK, True)
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(par.add_run("Agustus 2026"), size=9, color=GOLD, bold=True)


def add_toc(doc, guide):
    doc.add_page_break(); par = doc.add_paragraph(); set_font(par.add_run("Daftar Isi"), "Cambria", 23, GREEN, True)
    for idx, (title, _) in enumerate(guide["sections"], 1):
        t = doc.add_table(rows=1, cols=3); t.autofit = False; t.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells = t.rows[0].cells; widths = (550, 7600, 700)
        for cell, width in zip(cells, widths): set_cell_width(cell, width); set_cell_margins(cell, 30, 20, 30, 20)
        set_font(cells[0].paragraphs[0].add_run(f"{idx:02d}"), "Cambria", 11.5, GOLD, True)
        set_font(cells[1].paragraphs[0].add_run(title), "Cambria", 10.5, INK)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_font(cells[2].paragraphs[0].add_run(f"§{idx:02d}"), size=9, color=MUTED)


def add_section_heading(doc, index, title):
    doc.add_page_break()
    par = doc.add_paragraph(); set_font(par.add_run(f"BAGIAN {index:02d}"), "Cambria", 11, GOLD, True)
    par.paragraph_format.space_after = Pt(2)
    heading_style = next((item for item in doc.styles if item.name == "Heading 1"), None)
    title_par = doc.add_paragraph(style=heading_style); title_par.add_run(title)
    title_par.paragraph_format.space_after = Pt(10)
    p_pr = title_par._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "10"); bottom.set(qn("w:color"), GOLD); borders.append(bottom); p_pr.append(borders)


def add_blocks(doc, blocks):
    for block in blocks:
        kind = block["type"]
        if kind == "p": doc.add_paragraph(block["text"])
        elif kind == "link":
            par = doc.add_paragraph(); set_font(par.add_run(block["label"] + ": "), size=10.5, bold=True); add_hyperlink(par, block["url"], block["url"])
        elif kind == "h": doc.add_paragraph(block["text"], style=next((item for item in doc.styles if item.name == "Heading 2"), None))
        elif kind == "bullets":
            for item in block["items"]: add_bullet(doc, item)
        elif kind == "steps":
            add_table_separator(doc)
            t = doc.add_table(rows=len(block["items"]), cols=2); t.autofit = False
            grid_cols = t._tbl.tblGrid.findall(qn("w:gridCol"))
            for grid_col, width in zip(grid_cols, (430, 7900)):
                grid_col.set(qn("w:w"), str(width))
            for idx, (row, item) in enumerate(zip(t.rows, block["items"]), 1):
                n, body = row.cells; set_cell_width(n, 430); set_cell_width(body, 7900)
                set_cell_fill(n, GREEN_BRIGHT); set_cell_margins(n, 40, 50, 40, 50); set_cell_margins(body, 45, 120, 45, 40)
                n.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; body.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                n.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(n.paragraphs[0].add_run(str(idx)), size=9, color="FFFFFF", bold=True)
                set_font(body.paragraphs[0].add_run(item), size=10.5)
        elif kind == "note":
            fill = {"green": PALE_GREEN, "gold": PALE_GOLD, "red": PALE_RED}[block["tone"]]
            accent = {"green": GREEN_BRIGHT, "gold": GOLD, "red": "B5472A"}[block["tone"]]
            t = doc.add_table(rows=1, cols=1); cell = t.cell(0, 0); set_cell_fill(cell, fill); set_cell_margins(cell, 140, 180, 140, 180)
            tc_pr = cell._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders"); left = OxmlElement("w:left")
            left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "28"); left.set(qn("w:color"), accent); borders.append(left); tc_pr.append(borders)
            par = cell.paragraphs[0]; set_font(par.add_run(block["title"] + "\n"), "Cambria", 11, accent, True); set_font(par.add_run(block["text"]), size=10.5)
        elif kind == "table":
            cols = len(block["headers"]); t = doc.add_table(rows=1, cols=cols); t.autofit = False
            for i, head in enumerate(block["headers"]):
                cell = t.rows[0].cells[i]; set_cell_fill(cell, GREEN); set_cell_margins(cell); set_cell_width(cell, int(8500 / cols)); set_font(cell.paragraphs[0].add_run(head), size=9.5, color="FFFFFF", bold=True)
            set_repeat_table_header(t.rows[0])
            for row in block["rows"]:
                cells = t.add_row().cells
                for i, value in enumerate(row): set_cell_margins(cells[i]); set_cell_width(cells[i], int(8500 / cols)); set_font(cells[i].paragraphs[0].add_run(str(value)), size=9.5)
        elif kind == "image":
            path = ASSET_DIR / block["name"]
            if path.exists():
                par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                picture = par.add_run().add_picture(str(path), width=Inches(5.7))
                picture._inline.docPr.set("descr", block["caption"])
                picture._inline.docPr.set("title", block["caption"])
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_font(cap.add_run(block["caption"]), size=9, color=MUTED, italic=True)
        elif kind == "code":
            t = doc.add_table(rows=1, cols=1); cell = t.cell(0, 0); set_cell_fill(cell, "F1F3F2"); set_cell_margins(cell, 120, 150, 120, 150)
            set_font(cell.paragraphs[0].add_run(block["text"]), "Consolas", 8.5, GREEN)


def build_docx(guide):
    DOC_DIR.mkdir(parents=True, exist_ok=True); PUBLIC_FILES.mkdir(parents=True, exist_ok=True)
    target = DOC_DIR / f"{guide['filename']}.docx"
    shutil.copy2(REFERENCE, target)
    doc = Document(target); prepare_doc(doc, guide); add_cover(doc, guide); add_toc(doc, guide)
    for idx, (title, blocks_) in enumerate(guide["sections"], 1): add_section_heading(doc, idx, title); add_blocks(doc, blocks_)
    doc.core_properties.title = guide["title"] + " Portal Desa Rancajaya"
    doc.core_properties.subject = guide["description"]
    doc.core_properties.author = "KKN-T IPB 2026"
    doc.core_properties.keywords = "Portal Rancajaya, panduan website desa"
    doc.save(target)
    shutil.copy2(target, PUBLIC_FILES / target.name)
    return target


def block_html(block):
    kind = block["type"]
    if kind == "p": return f"<p>{html.escape(block['text'])}</p>"
    if kind == "link": return f"<p class='resource-link'><strong>{html.escape(block['label'])}:</strong> <a href='{html.escape(block['url'], quote=True)}' target='_blank' rel='noopener noreferrer'>{html.escape(block['url'])}</a></p>"
    if kind == "h": return f"<h3>{html.escape(block['text'])}</h3>"
    if kind in ("bullets", "steps"):
        tag = "ol" if kind == "steps" else "ul"; items = "".join(f"<li>{html.escape(x)}</li>" for x in block["items"]); return f"<{tag}>{items}</{tag}>"
    if kind == "note": return f"<aside class='note {block['tone']}'><strong>{html.escape(block['title'])}</strong><p>{html.escape(block['text'])}</p></aside>"
    if kind == "table":
        heads = "".join(f"<th>{html.escape(str(x))}</th>" for x in block["headers"])
        rows = "".join("<tr>" + "".join(f"<td>{html.escape(str(x))}</td>" for x in row) + "</tr>" for row in block["rows"])
        return f"<div class='table-wrap'><table><thead><tr>{heads}</tr></thead><tbody>{rows}</tbody></table></div>"
    if kind == "image": return f"<figure><img src='../panduan/assets/{html.escape(block['name'])}' alt='{html.escape(block['caption'])}'><figcaption>{html.escape(block['caption'])}</figcaption></figure>"
    if kind == "code": return f"<pre><code>{html.escape(block['text'])}</code></pre>"
    return ""


STYLE = """
:root{--green:#174c37;--green2:#166534;--gold:#b7791f;--ink:#252b28;--muted:#687386;--paper:#fbfaf7}*{box-sizing:border-box}html{scroll-behavior:smooth}body{margin:0;background:var(--paper);color:var(--ink);font:17px/1.65 system-ui,-apple-system,Segoe UI,sans-serif}.top{position:sticky;top:0;z-index:10;background:rgba(23,76,55,.97);color:#fff}.top-inner{max-width:1180px;margin:auto;padding:.9rem 1.25rem;display:flex;gap:1rem;align-items:center;justify-content:space-between}.top a{color:#fff;text-decoration:none;font-weight:750}.actions{display:flex;gap:.55rem;flex-wrap:wrap}.actions a{padding:.55rem .75rem;border:1px solid rgba(255,255,255,.35);border-radius:.5rem;font-size:.82rem}.layout{max-width:1180px;margin:auto;display:grid;grid-template-columns:270px minmax(0,1fr);gap:2.5rem;padding:2.5rem 1.25rem 5rem}.toc{position:sticky;top:5.5rem;align-self:start}.toc a{display:block;padding:.4rem .6rem;border-left:2px solid #d6ddd8;color:#365348;text-decoration:none;font-size:.86rem}.toc a:hover{border-color:var(--gold);color:var(--green)}main{min-width:0}.cover{padding:3rem 0 2rem;border-bottom:2px solid var(--gold);margin-bottom:2.6rem}.eyebrow{color:var(--gold);font-weight:800;letter-spacing:.14em}.cover h1{font:800 clamp(2.3rem,5vw,4.3rem)/1.06 Georgia,serif;color:var(--green);margin:.6rem 0}.cover h2{font:700 clamp(1.4rem,3vw,2.15rem)/1.2 Georgia,serif;color:var(--green);margin:0}.lede{font-size:1.12rem;color:var(--muted);max-width:65ch}.section{scroll-margin-top:6rem;margin:0 0 4rem}.section-label{color:var(--gold);font-weight:800}.section h2{font:800 clamp(1.8rem,3vw,2.5rem)/1.2 Georgia,serif;color:var(--green);padding-bottom:.55rem;border-bottom:2px solid var(--gold)}h3{font-family:Georgia,serif;color:var(--green)}li{margin:.45rem 0}.note{padding:1rem 1.2rem;margin:1.25rem 0;border-left:6px solid var(--green2);background:#e9f4ee}.note.gold{border-color:var(--gold);background:#fbf3dd}.note.red{border-color:#b5472a;background:#fbede7}.note p{margin:.25rem 0 0}.table-wrap{overflow:auto;margin:1.25rem 0}table{width:100%;border-collapse:collapse;background:#fff}th{background:var(--green);color:#fff;text-align:left}th,td{padding:.75rem;border:1px solid #d9cdb4;vertical-align:top}figure{margin:1.5rem 0;text-align:center}figure img{width:100%;border:1px solid #d9cdb4;background:#fff}figcaption{color:var(--muted);font-style:italic;font-size:.9rem}pre{overflow:auto;padding:1rem 1.1rem;background:#eff2f0;border-left:5px solid var(--green);color:#123d2d;font-size:.88rem}@media(max-width:850px){.layout{grid-template-columns:1fr}.toc{position:static}.top-inner{align-items:flex-start;flex-direction:column}.actions{width:100%}}
"""


def build_html(guide):
    out_dir = ROOT / "public" / "admin" / guide["slug"]
    out_dir.mkdir(parents=True, exist_ok=True)
    toc = "".join(f"<a href='#bagian-{i}'>{i:02d} {html.escape(title)}</a>" for i, (title, _) in enumerate(guide["sections"], 1))
    sections = "".join(f"<section class='section' id='bagian-{i}'><div class='section-label'>BAGIAN {i:02d}</div><h2>{html.escape(title)}</h2>{''.join(block_html(b) for b in blocks_)}</section>" for i, (title, blocks_) in enumerate(guide["sections"], 1))
    filename = guide["filename"]
    page = f"""<!doctype html><html lang='id'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><title>{html.escape(guide['title'])} - Portal Rancajaya</title><meta name='description' content='{html.escape(guide['description'])}'><style>{STYLE}</style></head><body><header class='top'><div class='top-inner'><a href='../index.html'>← Kembali ke Admin</a><div class='actions'><a href='../panduan/files/{filename}.pdf'>Unduh PDF</a><a href='../panduan/files/{filename}.docx'>Unduh Word</a></div></div></header><div class='layout'><nav class='toc' aria-label='Daftar isi'>{toc}</nav><main><div class='cover'><div class='eyebrow'>{guide['eyebrow']}</div><h1>{html.escape(guide['title'])}</h1><h2>{html.escape(guide['subtitle'])}</h2><p class='lede'>{html.escape(guide['description'])}</p><p><strong>Ditujukan untuk:</strong> {html.escape(guide['audience'])}<br><strong>Disusun oleh:</strong> KKN-T IPB 2026<br><strong>Versi:</strong> Agustus 2026</p></div>{sections}</main></div></body></html>"""
    page = page.replace("<style>", "<link rel='icon' href='../decap-cms-favicon.svg' type='image/svg+xml'><style>", 1)
    page = page.replace("<div class='actions'>", f"<div class='actions'><a href='{SUPPORT_URL}' target='_blank' rel='noopener noreferrer'>{SUPPORT_LABEL}</a>", 1)
    page = page.replace("<strong>Disusun oleh:</strong>", f"<strong>Bantuan teknis:</strong> <a href='{SUPPORT_URL}' target='_blank' rel='noopener noreferrer'>Daffa via WhatsApp</a><br><strong>Disusun oleh:</strong>", 1)
    (out_dir / "index.html").write_text(page, encoding="utf-8")


if __name__ == "__main__":
    for guide in (OPERATIONAL, DEVELOPMENT):
        build_docx(guide)
        build_html(guide)
